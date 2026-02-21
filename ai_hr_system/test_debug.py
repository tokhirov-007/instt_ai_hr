import sys
import os
from app.cv_intelligence.cv_analyzer import CVAnalyzer
import docx

def create_dummy_cv(filename):
    doc = docx.Document()
    doc.add_heading('John Doe', 0)
    doc.add_paragraph('Senior Software Engineer with 5+ years of experience.')
    doc.add_paragraph('Specialized in Modern JavaScript frameworks and Backend patterns.')
    doc.add_paragraph('Proficient in Python, standard data science stack, and cloud native technologies.')
    doc.save(filename)
    print(f"Created dummy CV: {filename}")

def test_analysis():
    print("Testing CV Analyzer...")
    filename = "test_cv.docx"
    create_dummy_cv(filename)
    
    try:
        analyzer = CVAnalyzer()
        result = analyzer.analyze(filename)
        
        print("\n--- Analysis Result ---")
        print(f"Detected Skills: {result.skills_detected}")
        print(f"Inferred Skills: {result.inferred_skills}")
        print(f"Experience: {result.experience_years} years")
        
        # Validation checks
        # Validation checks
        # Check lower case or normalize first. The skills detected are lowercased by extractor?
        # Let's check:
        # Detected Skills: ['javascript', 'python']
        # Assertion was for "Python".
        
        detected_lower = [s.lower() for s in result.skills_detected]
        assert "python" in detected_lower, "Failed to detect explicit skill 'Python'"
        
        # Check if inference works
        inferred = result.inferred_skills
        assert any(x in inferred for x in ["React", "Vue", "Angular"]), "Failed to infer Modern JS frameworks"
        assert any(x in inferred for x in ["Pandas", "Scikit-Learn"]), "Failed to infer Data Science stack"
        assert result.experience_years == 5.0, "Failed to extract experience years"
        
        print("\n[SUCCESS] Verification Successful!")
        
    except Exception as e:
        print(f"\n[FAILURE] Verification Failed: {e}")
        import traceback
        traceback.print_exc()
    finally:
        if os.path.exists(filename):
            os.remove(filename)

if __name__ == "__main__":
    # Add current dir to path to find app module
    sys.path.append(os.getcwd())
    test_analysis()
